import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_docx_workbook():
    doc = Document()
    
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.85)
        s.right_margin = Inches(0.85)

    COLOR_GOLD = RGBColor(180, 135, 30)
    COLOR_NAVY = RGBColor(14, 18, 27)
    COLOR_MUTED = RGBColor(90, 100, 115)
    COLOR_CORAL = RGBColor(233, 69, 96)

    # PORTADA / ENCABEZADO
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_header.add_run("MÉTODO MODO LÍDER • WORKBOOK DE APOYO\n")
    r_sub.font.name = "Montserrat"
    r_sub.font.size = Pt(10)
    r_sub.font.bold = True
    r_sub.font.color.rgb = COLOR_GOLD

    r_title = p_header.add_run("CAPITULO 1 EMPRENDER CONSCIENTEMENTE\n")
    r_title.font.name = "Montserrat"
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_NAVY

    doc.add_paragraph()

    # BLOQUE 1: ¿QUÉ DIFERENCIA...?
    p1 = doc.add_paragraph()
    r = p1.add_run("¿Qué diferencia a un emprendedor consciente de otro que no lo es?\n")
    r.font.name = "Montserrat"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = COLOR_NAVY

    r2 = p1.add_run("La diferencia clave podría ser:\n\n")
    r2.font.name = "Montserrat"
    r2.font.size = Pt(10.5)

    t_box1 = doc.add_table(rows=2, cols=1)
    c1 = t_box1.cell(0, 0)
    set_cell_background(c1, "FFF5F6")
    set_cell_margins(c1, 140, 140, 180, 180)
    p_c1 = c1.paragraphs[0]
    r_q1 = p_c1.add_run("Un emprendedor no consciente pregunta: “¿Cómo logro que mi negocio funcione?”")
    r_q1.font.name = "Montserrat"
    r_q1.font.size = Pt(10)
    r_q1.font.bold = True
    r_q1.font.color.rgb = COLOR_CORAL

    c2 = t_box1.cell(1, 0)
    set_cell_background(c2, "FEFBF2")
    set_cell_margins(c2, 140, 140, 180, 180)
    p_c2 = c2.paragraphs[0]
    r_q2 = p_c2.add_run("Un emprendedor consciente pregunta: “¿Qué negocio quiero construir y en quién me quiero convertir mientras lo construyo?” y luego se pregunta otras cosas, como con qué estilo de clientes me gustaría trabajar . ¿Cómo logro comunicarme con ellos? ¿Cómo conecto con sus intereses para crear conexión real? Para luego pensar, ¿qué cosas les gustan y les disgustan a mis clientes objetivos? ¿Qué odian, qué temen, que necesitan y de cuáles de sus necesidades voy a ocuparme de resolver?")
    r_q2.font.name = "Montserrat"
    r_q2.font.size = Pt(10)
    r_q2.font.bold = True

    doc.add_paragraph()

    # BLOQUE 2: 🌱 La diferencia no está en el negocio...
    p2 = doc.add_paragraph()
    r = p2.add_run("🌱 La diferencia no está en el negocio, sino en la manera de pensar y de emprender tu negocio\n\n")
    r.font.name = "Montserrat"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = COLOR_NAVY

    r = p2.add_run("Un emprendedor consciente también quiere ganar dinero, crecer, vender y tener éxito. No renuncia a la ambición.\n\nLa diferencia es que no separa el éxito económico del impacto que genera, de sus valores y de la persona que está siendo.\n\nPor otro lado un emprendedor consciente sabe que para poder lograr lo que desea necesariamente necesita trascender sus miedos, desactivar sus limitaciones mentales y/o reconocer y resolver las heridas que bloquean su camino y limitaciones hacia su éxito.\n")
    r.font.name = "Montserrat"
    r.font.size = Pt(10)

    # CITA
    t_quote = doc.add_table(rows=1, cols=1)
    cq = t_quote.cell(0, 0)
    set_cell_background(cq, "FEF9EB")
    set_cell_margins(cq, 160, 160, 200, 200)
    pq = cq.paragraphs[0]
    pq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rq = pq.add_run("“Ser un emprendedor consciente no significa emprender sin ambición. Significa tener ambición sin perderte a vos misma en el camino.”")
    rq.font.name = "Montserrat"
    rq.font.size = Pt(10.5)
    rq.font.bold = True
    rq.font.color.rgb = COLOR_GOLD

    doc.add_paragraph()

    # TABLA COMPARATIVA
    matriz_data = [
        ("Busca resultados rápidos", "Construye relaciones y conexión hacia resultados sostenibles"),
        ("Decide principalmente por dinero", "Decide considerando propósito + impacto hacia el dinero"),
        ("Vende lo que puede", "Busca resolver problemas reales de los clientes que decidió atender"),
        ("Compite permanentemente", "Puede colaborar y crear redes desde el respeto, los límites y los acuerdos"),
        ("Trabaja hasta agotarse", "Aprende a gestionar su bienestar, su energía y recursos"),
        ("Busca demostrar que puede", "Busca construir algo que tenga sentido y que genere dinero"),
        ("Se identifica con el resultado", "Aprende del resultado y sigue mejorando o escalando"),
        ("Tiene miedo de equivocarse", "Usa el error como información y crecimiento"),
        ("Entiende que puede crecer a cualquier costo porque así es como se gana dinero", "Define qué costo no está dispuesto a pagar con sabiduría y en coherencia con sus valores"),
        ("Pregunta “¿cuánto gano?”", "Pregunta “¿Qué genero y qué construyo y cuánto gano?”")
    ]

    t_mat = doc.add_table(rows=11, cols=2)
    t_mat.alignment = WD_TABLE_ALIGNMENT.CENTER

    h0 = t_mat.cell(0, 0)
    h0.text = "Emprendedor tradicional / no consciente"
    set_cell_background(h0, "FFF0F3")
    h0.paragraphs[0].runs[0].font.name = "Montserrat"
    h0.paragraphs[0].runs[0].font.bold = True
    h0.paragraphs[0].runs[0].font.color.rgb = COLOR_CORAL
    h0.paragraphs[0].runs[0].font.size = Pt(9.5)

    h1 = t_mat.cell(0, 1)
    h1.text = "Emprendedor consciente"
    set_cell_background(h1, "FEF9EB")
    h1.paragraphs[0].runs[0].font.name = "Montserrat"
    h1.paragraphs[0].runs[0].font.bold = True
    h1.paragraphs[0].runs[0].font.color.rgb = COLOR_GOLD
    h1.paragraphs[0].runs[0].font.size = Pt(9.5)

    for i, (col0, col1) in enumerate(matriz_data, start=1):
        bg = "FFFFFF" if i % 2 == 0 else "F9FAFC"
        c0 = t_mat.cell(i, 0)
        c0.text = col0
        set_cell_background(c0, bg)
        c0.paragraphs[0].runs[0].font.name = "Montserrat"
        c0.paragraphs[0].runs[0].font.size = Pt(9)

        c1 = t_mat.cell(i, 1)
        c1.text = col1
        set_cell_background(c1, bg)
        c1.paragraphs[0].runs[0].font.name = "Montserrat"
        c1.paragraphs[0].runs[0].font.bold = True
        c1.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # PREGUNTAS INCÓMODAS
    p3 = doc.add_paragraph()
    r = p3.add_run("🔥 Y después haría una pregunta incómoda\n")
    r.font.name = "Montserrat"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = COLOR_NAVY

    r = p3.add_run("Les preguntaría:\n\n")
    r.font.name = "Montserrat"
    r.font.size = Pt(10.5)

    t_inc = doc.add_table(rows=3, cols=1)
    inc_items = [
        "¿Qué estás dispuesta a hacer para que tu emprendimiento tenga éxito? Aquí es donde se crea la mentalidad que necesitás para llegar lejos",
        "“¿Y qué NO estás dispuesta a hacer?” Aquí es donde aparecen tus valores, tus límites, tus criterios de decisión básicos",
        "¿Cuánto necesitás cobrar por tus productos y/o servicios para que tu propia economía + tu negocio funcione?"
    ]
    for idx, txt in enumerate(inc_items):
        c = t_inc.cell(idx, 0)
        set_cell_background(c, "FDFDFD")
        set_cell_margins(c, 160, 160, 200, 200)
        p = c.paragraphs[0]
        r = p.add_run(txt + "\n\n\n\n")
        r.font.name = "Montserrat"
        r.font.size = Pt(10)
        r.font.bold = True

    doc.add_paragraph()

    # FÓRMULA
    p4 = doc.add_paragraph()
    r = p4.add_run("EMPRENDER CONSCIENTEMENTE = PROPÓSITO + IDENTIDAD + MENTALIDAD + ACCIÓN CON ESTRATEGIA = IMPACTO + DINERO QUE FLUYE HACIA TU NEGOCIO\n\n")
    r.font.name = "Montserrat"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = COLOR_GOLD

    formula_bullets = [
        ("• Propósito:", "¿para qué?"),
        ("• Identidad:", "¿quién soy y qué quiero construir?"),
        ("• Mentalidad:", "¿qué necesito transformar?"),
        ("• Acción:", "¿qué hago?"),
        ("• Estrategia:", "¿cómo lo hago inteligentemente?"),
        ("• Impacto:", "¿qué genero?"),
        ("• Dinero:", "¿cómo genero dinero para que mi sustento personal sea sostenible?")
    ]
    for b_l, b_d in formula_bullets:
        pb = doc.add_paragraph()
        rb1 = pb.add_run(f"{b_l} ")
        rb1.font.name = "Montserrat"
        rb1.font.bold = True
        rb1.font.size = Pt(10)
        rb2 = pb.add_run(b_d)
        rb2.font.name = "Montserrat"
        rb2.font.size = Pt(10)

    doc.add_paragraph()

    # TEXTO CONCIENCIA EMPRESARIAL
    p5 = doc.add_paragraph()
    r = p5.add_run("Un emprendedor consciente sabe quién es, tiene una dirección (propósito) y cada día toma decisiones y acciona en función de sus valores y transformando su mentalidad hacia  lo que necesita hacer para llegar un día a la vez hacia crear y sostener un negocio de impacto que le permita vivir con el estilo de vida que desea.\n\nLa conciencia empresarial es aprender a tomar mejores decisiones con los elementos que se tienen disponibles pero sin perder de vista los valores los limites la conexion con el cliente y lo que deseo para mi y mi negocio.\n\nEstas decisiones puestas en acción a diario crean una MARCA PERSONAL QUE SE ENTIENDE Y FUNCIONA PARA GENERAR DINERO\n")
    r.font.name = "Montserrat"
    r.font.size = Pt(10)

    # CIERRE
    t_end = doc.add_table(rows=1, cols=1)
    ce = t_end.cell(0, 0)
    set_cell_background(ce, "FEF9EB")
    set_cell_margins(ce, 180, 180, 220, 220)
    pe = ce.paragraphs[0]
    pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    re = pe.add_run("“Vamos a aprender a construir un negocio sin desconectarnos de nosotras mismas, pero tampoco vamos a romantizarlo porque emprender es crear un negocio que genere renta que te sustente a vos y tu estilo de vida . Vamos a hablar de propósito, dinero, ventas, de mentalidad y autoliderazgo, pero también de estrategias claves y acción consciente para generar negocios con valor de marca e impacto real que atrae dinero.”")
    re.font.name = "Montserrat"
    re.font.size = Pt(10.5)
    re.font.bold = True

    out_docx = r"c:\Users\ecasa\Documents\MODO LIDER\METODOMODOLIDER\clase1_emprender_conscientemente\Workbook_Clase_1_Emprender_Conscientemente.docx"
    doc.save(out_docx)
    print(f"DOCX saved with verbatim text at: {out_docx}")

if __name__ == "__main__":
    create_docx_workbook()
